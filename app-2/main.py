import os
import json
import threading
from pathlib import Path
import customtkinter as ctk
from tkinter import filedialog, messagebox
from PIL import Image
import pandas as pd
from docx import Document
import pdfplumber
import pythoncom
import win32com.client as win32

# Настройка темы
ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("dark-blue")


class LanguageManager:
    def __init__(self):
        self.languages = {
            'en': {
                'title': 'Quick Converter',
                'select_file': 'Select File',
                'convert': 'Convert',
                'output_format': 'Output Format',
                'output_folder': 'Output Folder',
                'browse': 'Browse',
                'converting': 'Converting...',
                'success': 'Conversion completed successfully!',
                'error': 'Error during conversion',
                'no_file': 'Please select a file',
                'no_format': 'Please select output format',
                'no_output_folder': 'Please select output folder',
                'supported_formats': 'Supported Formats',
                'language': 'Language',
                'settings': 'Settings',
                'file_size': 'File Size',
                'status': 'Status',
                'ready': 'Ready',
                'progress': 'Progress',
                'cancel': 'Cancel',
                'image_formats': 'Image Formats',
                'document_formats': 'Document Formats',
                'data_formats': 'Data Formats',
                'all_files': 'All Files',
                'unsupported_format': 'Unsupported file format'
            },
            'ru': {
                'title': 'Быстрый Конвертер',
                'select_file': 'Выбрать Файл',
                'convert': 'Конвертировать',
                'output_format': 'Выходной Формат',
                'output_folder': 'Папка назначения',
                'browse': 'Обзор',
                'converting': 'Конвертация...',
                'success': 'Конвертация успешно завершена!',
                'error': 'Ошибка при конвертации',
                'no_file': 'Пожалуйста, выберите файл',
                'no_format': 'Пожалуйста, выберите выходной формат',
                'no_output_folder': 'Пожалуйста, выберите папку назначения',
                'supported_formats': 'Поддерживаемые Форматы',
                'language': 'Язык',
                'settings': 'Настройки',
                'file_size': 'Размер файла',
                'status': 'Статус',
                'ready': 'Готов',
                'progress': 'Прогресс',
                'cancel': 'Отмена',
                'image_formats': 'Форматы изображений',
                'document_formats': 'Форматы документов',
                'data_formats': 'Форматы данных',
                'all_files': 'Все файлы',
                'unsupported_format': 'Неподдерживаемый формат файла'
            },
            'zh': {
                'title': '快速转换器',
                'select_file': '选择文件',
                'convert': '转换',
                'output_format': '输出格式',
                'output_folder': '输出文件夹',
                'browse': '浏览',
                'converting': '转换中...',
                'success': '转换成功完成！',
                'error': '转换过程中出错',
                'no_file': '请选择文件',
                'no_format': '请选择输出格式',
                'no_output_folder': '请选择输出文件夹',
                'supported_formats': '支持的格式',
                'language': '语言',
                'settings': '设置',
                'file_size': '文件大小',
                'status': '状态',
                'ready': '准备就绪',
                'progress': '进度',
                'cancel': '取消',
                'image_formats': '图像格式',
                'document_formats': '文档格式',
                'data_formats': '数据格式',
                'all_files': '所有文件',
                'unsupported_format': '不支持的文件格式'
            },
            'es': {
                'title': 'Convertidor Rápido',
                'select_file': 'Seleccionar Archivo',
                'convert': 'Convertir',
                'output_format': 'Formato de Salida',
                'output_folder': 'Carpeta de Salida',
                'browse': 'Examinar',
                'converting': 'Convirtiendo...',
                'success': '¡Conversión completada con éxito!',
                'error': 'Error durante la conversión',
                'no_file': 'Por favor seleccione un archivo',
                'no_format': 'Por favor seleccione el formato de salida',
                'no_output_folder': 'Por favor seleccione la carpeta de salida',
                'supported_formats': 'Formatos Soportados',
                'language': 'Idioma',
                'settings': 'Configuración',
                'file_size': 'Tamaño del Archivo',
                'status': 'Estado',
                'ready': 'Listo',
                'progress': 'Progreso',
                'cancel': 'Cancelar',
                'image_formats': 'Formatos de Imagen',
                'document_formats': 'Formatos de Documento',
                'data_formats': 'Formatos de Datos',
                'all_files': 'Todos los archivos',
                'unsupported_format': 'Formato de archivo no compatible'
            }
        }
        self.current_lang = 'en'
        self.load_settings()

    def load_settings(self):
        try:
            if os.path.exists('settings.json'):
                with open('settings.json', 'r', encoding='utf-8') as f:
                    settings = json.load(f)
                    lang = settings.get('language', 'en')
                    # Проверяем, что язык существует в нашем словаре
                    if lang in self.languages:
                        self.current_lang = lang
                    else:
                        self.current_lang = 'en'
        except Exception as e:
            print(f"Error loading settings: {e}")
            self.current_lang = 'en'

    def save_settings(self):
        try:
            settings = {'language': self.current_lang}
            with open('settings.json', 'w', encoding='utf-8') as f:
                json.dump(settings, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Error saving settings: {e}")

    def set_language(self, lang):
        if lang in self.languages:
            self.current_lang = lang
            self.save_settings()
        else:
            print(f"Language {lang} not found!")

    def get_text(self, key):
        return self.languages.get(self.current_lang, self.languages['en']).get(key, key)


class FileConverter:
    def __init__(self):
        self.supported_formats = {
            'image': {
                'extensions': ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp', '.gif'],
                'convert_to': ['.jpg', '.png', '.bmp', '.tiff', '.webp']
            },
            'document': {
                'extensions': ['.docx', '.doc', '.pdf', '.txt', '.rtf'],
                'convert_to': ['.docx', '.pdf', '.txt']
            },
            'data': {
                'extensions': ['.xlsx', '.xls', '.csv', '.json'],
                'convert_to': ['.xlsx', '.csv', '.json']
            }
        }

    def get_file_type(self, file_path):
        ext = Path(file_path).suffix.lower()
        for file_type, info in self.supported_formats.items():
            if ext in info['extensions']:
                return file_type
        return None

    def convert_image(self, input_path, output_path, lang_manager):
        try:
            with Image.open(input_path) as img:
                # Конвертируем в RGB если нужно (для форматов без альфа-канала)
                if img.mode in ('RGBA', 'LA') and Path(output_path).suffix.lower() in ['.jpg', '.jpeg']:
                    img = img.convert('RGB')
                img.save(output_path)
            return True
        except Exception as e:
            messagebox.showerror(lang_manager.get_text('error'), str(e))
            return False

    def convert_document(self, input_path, output_path, lang_manager):
        try:
            input_ext = Path(input_path).suffix.lower()
            output_ext = Path(output_path).suffix.lower()

            if input_ext == '.docx' and output_ext == '.pdf':
                # Конвертация DOCX в PDF
                pythoncom.CoInitialize()
                try:
                    word = win32.DispatchEx("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(input_path)
                    doc.SaveAs(output_path, FileFormat=17)  # 17 = PDF format
                    doc.Close()
                    word.Quit()
                finally:
                    pythoncom.CoUninitialize()

            elif input_ext == '.pdf' and output_ext == '.docx':
                # Конвертация PDF в DOCX (базовая)
                doc = Document()
                with pdfplumber.open(input_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            doc.add_paragraph(text)
                doc.save(output_path)

            elif input_ext in ['.docx', '.doc'] and output_ext == '.txt':
                # DOCX/DOC в TXT
                if input_ext == '.docx':
                    doc = Document(input_path)
                    text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                else:
                    # Для .doc файлов
                    pythoncom.CoInitialize()
                    try:
                        word = win32.DispatchEx("Word.Application")
                        word.Visible = False
                        doc = word.Documents.Open(input_path)
                        text = doc.Content.Text
                        doc.Close()
                        word.Quit()
                    finally:
                        pythoncom.CoUninitialize()

                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)

            elif input_ext == '.txt' and output_ext == '.docx':
                # TXT в DOCX
                doc = Document()
                with open(input_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                doc.add_paragraph(text)
                doc.save(output_path)

            return True
        except Exception as e:
            messagebox.showerror(lang_manager.get_text('error'), str(e))
            return False

    def convert_data(self, input_path, output_path, lang_manager):
        try:
            input_ext = Path(input_path).suffix.lower()
            output_ext = Path(output_path).suffix.lower()

            if input_ext in ['.xlsx', '.xls']:
                df = pd.read_excel(input_path)
            elif input_ext == '.csv':
                df = pd.read_csv(input_path)
            elif input_ext == '.json':
                df = pd.read_json(input_path)
            else:
                return False

            if output_ext == '.xlsx':
                df.to_excel(output_path, index=False)
            elif output_ext == '.csv':
                df.to_csv(output_path, index=False)
            elif output_ext == '.json':
                df.to_json(output_path, orient='records', indent=2, force_ascii=False)
            else:
                return False

            return True
        except Exception as e:
            messagebox.showerror(lang_manager.get_text('error'), str(e))
            return False


class QuickConverterApp:
    def __init__(self):
        self.root = ctk.CTk()
        self.lang_manager = LanguageManager()
        self.converter = FileConverter()
        self.current_file = None
        self.conversion_thread = None
        self.setup_ui()
        self.update_ui_text()  # Первоначальное обновление текста

    def setup_ui(self):
        self.root.title(self.lang_manager.get_text('title'))
        self.root.geometry('800x900')
        self.root.resizable(False, False)
        self.root.minsize(700, 500)

        # Создание основного фрейма
        self.main_frame = ctk.CTkFrame(self.root)
        self.main_frame.pack(fill='both', expand=True, padx=20, pady=20)

        # Заголовок
        self.title_label = ctk.CTkLabel(
            self.main_frame,
            text="",  # Будет установлено в update_ui_text
            font=ctk.CTkFont(size=24, weight='bold')
        )
        self.title_label.pack(pady=(0, 30))

        # Фрейм выбора файла
        self.file_frame = ctk.CTkFrame(self.main_frame)
        self.file_frame.pack(fill='x', pady=(0, 20))

        self.file_label = ctk.CTkLabel(
            self.file_frame,
            text="",  # Будет установлено в update_ui_text
            font=ctk.CTkFont(size=14, weight='bold')
        )
        self.file_label.pack(anchor='w', padx=20, pady=(20, 10))

        self.file_selection_frame = ctk.CTkFrame(self.file_frame, fg_color='transparent')
        self.file_selection_frame.pack(fill='x', padx=20, pady=(0, 20))

        self.file_path_entry = ctk.CTkEntry(
            self.file_selection_frame,
            placeholder_text="",  # Будет установлено в update_ui_text
            height=40
        )
        self.file_path_entry.pack(side='left', fill='x', expand=True, padx=(0, 10))

        self.browse_button = ctk.CTkButton(
            self.file_selection_frame,
            text="",  # Будет установлено в update_ui_text
            command=self.browse_file,
            width=100,
            height=40
        )
        self.browse_button.pack(side='right')

        # Информация о файле
        self.file_info_frame = ctk.CTkFrame(self.main_frame)
        self.file_info_frame.pack(fill='x', pady=(0, 20))

        self.file_size_label = ctk.CTkLabel(
            self.file_info_frame,
            text="",  # Будет установлено в update_ui_text
            font=ctk.CTkFont(size=12)
        )
        self.file_size_label.pack(anchor='w', padx=20, pady=10)

        # Фрейм форматов
        self.format_frame = ctk.CTkFrame(self.main_frame)
        self.format_frame.pack(fill='x', pady=(0, 20))

        self.format_label = ctk.CTkLabel(
            self.format_frame,
            text="",  # Будет установлено в update_ui_text
            font=ctk.CTkFont(size=14, weight='bold')
        )
        self.format_label.pack(anchor='w', padx=20, pady=(20, 10))

        # Выбор формата
        self.format_var = ctk.StringVar()
        self.format_combobox = ctk.CTkComboBox(
            self.format_frame,
            values=[],
            variable=self.format_var,
            state='readonly',
            height=40
        )
        self.format_combobox.pack(fill='x', padx=20, pady=(0, 20))

        # Фрейм папки назначения
        self.output_frame = ctk.CTkFrame(self.main_frame)
        self.output_frame.pack(fill='x', pady=(0, 20))

        self.output_label = ctk.CTkLabel(
            self.output_frame,
            text="",  # Будет установлено в update_ui_text
            font=ctk.CTkFont(size=14, weight='bold')
        )
        self.output_label.pack(anchor='w', padx=20, pady=(20, 10))

        self.output_selection_frame = ctk.CTkFrame(self.output_frame, fg_color='transparent')
        self.output_selection_frame.pack(fill='x', padx=20, pady=(0, 20))

        self.output_path_entry = ctk.CTkEntry(
            self.output_selection_frame,
            placeholder_text="",  # Будет установлено в update_ui_text
            height=40
        )
        self.output_path_entry.pack(side='left', fill='x', expand=True, padx=(0, 10))
        self.output_path_entry.insert(0, str(Path.home() / 'Downloads'))

        self.output_browse_button = ctk.CTkButton(
            self.output_selection_frame,
            text="",  # Будет установлено в update_ui_text
            command=self.browse_output_folder,
            width=100,
            height=40
        )
        self.output_browse_button.pack(side='right')

        # Прогресс бар
        self.progress_frame = ctk.CTkFrame(self.main_frame)
        self.progress_frame.pack(fill='x', pady=(0, 20))

        self.progress_label = ctk.CTkLabel(
            self.progress_frame,
            text="",  # Будет установлено в update_ui_text
            font=ctk.CTkFont(size=12, weight='bold')
        )
        self.progress_label.pack(anchor='w', padx=20, pady=(20, 10))

        self.progress_bar = ctk.CTkProgressBar(self.progress_frame)
        self.progress_bar.pack(fill='x', padx=20, pady=(0, 10))
        self.progress_bar.set(0)

        self.status_label = ctk.CTkLabel(
            self.progress_frame,
            text="",  # Будет установлено в update_ui_text
            font=ctk.CTkFont(size=12)
        )
        self.status_label.pack(anchor='w', padx=20, pady=(0, 20))

        # Кнопки
        self.button_frame = ctk.CTkFrame(self.main_frame, fg_color='transparent')
        self.button_frame.pack(fill='x', pady=(0, 20))

        self.convert_button = ctk.CTkButton(
            self.button_frame,
            text="",  # Будет установлено в update_ui_text
            command=self.start_conversion,
            height=45,
            font=ctk.CTkFont(size=16, weight='bold')
        )
        self.convert_button.pack(side='left', padx=(0, 10))

        self.cancel_button = ctk.CTkButton(
            self.button_frame,
            text="",  # Будет установлено в update_ui_text
            command=self.cancel_conversion,
            height=45,
            fg_color='gray',
            hover_color='dark gray',
            state='disabled'
        )
        self.cancel_button.pack(side='left')

        # Настройки
        self.settings_frame = ctk.CTkFrame(self.main_frame)
        self.settings_frame.pack(fill='x')

        self.settings_label = ctk.CTkLabel(
            self.settings_frame,
            text="",  # Будет установлено в update_ui_text
            font=ctk.CTkFont(size=14, weight='bold')
        )
        self.settings_label.pack(anchor='w', padx=20, pady=(20, 10))

        self.language_combobox = ctk.CTkComboBox(
            self.settings_frame,
            values=['English', 'Русский', '中文', 'Español'],
            command=self.change_language,
            state='readonly',
            height=40
        )
        self.language_combobox.pack(fill='x', padx=20, pady=(0, 20))
        self.set_language_combobox_value()

    def update_ui_text(self):
        """Полное обновление всего текста в интерфейсе"""
        # Основные элементы
        self.root.title(self.lang_manager.get_text('title'))
        self.title_label.configure(text=self.lang_manager.get_text('title'))
        self.file_label.configure(text=self.lang_manager.get_text('select_file'))
        self.file_path_entry.configure(placeholder_text=self.lang_manager.get_text('select_file'))
        self.browse_button.configure(text=self.lang_manager.get_text('browse'))
        self.format_label.configure(text=self.lang_manager.get_text('output_format'))
        self.output_label.configure(text=self.lang_manager.get_text('output_folder'))
        self.output_path_entry.configure(placeholder_text=self.lang_manager.get_text('output_folder'))
        self.output_browse_button.configure(text=self.lang_manager.get_text('browse'))
        self.progress_label.configure(text=f"{self.lang_manager.get_text('progress')}:")
        self.convert_button.configure(text=self.lang_manager.get_text('convert'))
        self.cancel_button.configure(text=self.lang_manager.get_text('cancel'))
        self.settings_label.configure(text=self.lang_manager.get_text('settings'))

        # Обновляем информацию о файле
        if self.current_file:
            self.update_file_info(self.current_file)
        else:
            self.file_size_label.configure(text=f"{self.lang_manager.get_text('file_size')}: -")

        # Обновляем статус
        self.status_label.configure(text=self.lang_manager.get_text('ready'))

        # Обновляем форматы файлов для диалога выбора
        self.update_file_dialog_formats()

    def update_file_dialog_formats(self):
        """Обновляет форматы файлов в диалоге выбора"""
        self.file_dialog_formats = [
            (self.lang_manager.get_text('all_files'), '*.*'),
            (self.lang_manager.get_text('image_formats'), '*.jpg *.jpeg *.png *.bmp *.tiff *.webp *.gif'),
            (self.lang_manager.get_text('document_formats'), '*.docx *.doc *.pdf *.txt *.rtf'),
            (self.lang_manager.get_text('data_formats'), '*.xlsx *.xls *.csv *.json')
        ]

    def set_language_combobox_value(self):
        lang_map = {'en': 'English', 'ru': 'Русский', 'zh': '中文', 'es': 'Español'}
        self.language_combobox.set(lang_map.get(self.lang_manager.current_lang, 'English'))

    def browse_file(self):
        filename = filedialog.askopenfilename(filetypes=self.file_dialog_formats)
        if filename:
            self.current_file = filename
            self.file_path_entry.delete(0, 'end')
            self.file_path_entry.insert(0, filename)
            self.update_file_info(filename)
            self.update_format_options(filename)

    def browse_output_folder(self):
        folder = filedialog.askdirectory()
        if folder:
            self.output_path_entry.delete(0, 'end')
            self.output_path_entry.insert(0, folder)

    def update_file_info(self, file_path):
        try:
            size = os.path.getsize(file_path)
            size_kb = size / 1024
            if size_kb < 1024:
                size_str = f"{size_kb:.1f} KB"
            else:
                size_mb = size_kb / 1024
                size_str = f"{size_mb:.1f} MB"

            self.file_size_label.configure(
                text=f"{self.lang_manager.get_text('file_size')}: {size_str}"
            )
        except:
            self.file_size_label.configure(
                text=f"{self.lang_manager.get_text('file_size')}: -"
            )

    def update_format_options(self, file_path):
        file_type = self.converter.get_file_type(file_path)
        if file_type:
            formats = self.converter.supported_formats[file_type]['convert_to']
            format_names = [f.upper() + ' ' + self.get_format_description(f) for f in formats]
            self.format_combobox.configure(values=format_names)
            if formats:
                self.format_combobox.set(format_names[0])
        else:
            self.format_combobox.configure(values=[])
            self.format_var.set('')

    def get_format_description(self, ext):
        descriptions = {
            '.jpg': self.lang_manager.get_text('image_formats').split(' ')[0] + ' JPEG',
            '.png': self.lang_manager.get_text('image_formats').split(' ')[0] + ' PNG',
            '.bmp': self.lang_manager.get_text('image_formats').split(' ')[0] + ' BMP',
            '.tiff': self.lang_manager.get_text('image_formats').split(' ')[0] + ' TIFF',
            '.webp': self.lang_manager.get_text('image_formats').split(' ')[0] + ' WebP',
            '.docx': self.lang_manager.get_text('document_formats').split(' ')[0] + ' DOCX',
            '.pdf': self.lang_manager.get_text('document_formats').split(' ')[0] + ' PDF',
            '.txt': self.lang_manager.get_text('document_formats').split(' ')[0] + ' TXT',
            '.xlsx': self.lang_manager.get_text('data_formats').split(' ')[0] + ' XLSX',
            '.csv': self.lang_manager.get_text('data_formats').split(' ')[0] + ' CSV',
            '.json': self.lang_manager.get_text('data_formats').split(' ')[0] + ' JSON'
        }
        return descriptions.get(ext, ext.upper() + ' File')

    def change_language(self, choice):
        lang_map = {'English': 'en', 'Русский': 'ru', '中文': 'zh', 'Español': 'es'}
        new_lang = lang_map.get(choice, 'en')
        self.lang_manager.set_language(new_lang)
        self.update_ui_text()

    def start_conversion(self):
        if not self.current_file:
            messagebox.showwarning(
                self.lang_manager.get_text('error'),
                self.lang_manager.get_text('no_file')
            )
            return

        if not self.format_var.get():
            messagebox.showwarning(
                self.lang_manager.get_text('error'),
                self.lang_manager.get_text('no_format')
            )
            return

        output_folder = self.output_path_entry.get()
        if not output_folder or not os.path.exists(output_folder):
            messagebox.showwarning(
                self.lang_manager.get_text('error'),
                self.lang_manager.get_text('no_output_folder')
            )
            return

        # Получаем выбранный формат
        selected_format = self.format_combobox.get().split(' ')[0].lower()
        if not selected_format.startswith('.'):
            selected_format = '.' + selected_format

        # Создаем путь для выходного файла
        input_path = Path(self.current_file)
        output_path = Path(output_folder) / f"{input_path.stem}_converted{selected_format}"

        # Запускаем конвертацию в отдельном потоке
        self.conversion_thread = threading.Thread(
            target=self.convert_file,
            args=(str(input_path), str(output_path), selected_format)
        )
        self.conversion_thread.daemon = True
        self.conversion_thread.start()

        # Обновляем UI
        self.convert_button.configure(state='disabled')
        self.cancel_button.configure(state='normal')
        self.progress_bar.set(0.5)
        self.status_label.configure(text=self.lang_manager.get_text('converting'))

    def convert_file(self, input_path, output_path, output_format):
        try:
            file_type = self.converter.get_file_type(input_path)
            success = False

            if file_type == 'image':
                success = self.converter.convert_image(input_path, output_path, self.lang_manager)
            elif file_type == 'document':
                success = self.converter.convert_document(input_path, output_path, self.lang_manager)
            elif file_type == 'data':
                success = self.converter.convert_data(input_path, output_path, self.lang_manager)
            else:
                self.root.after(0, lambda: messagebox.showerror(
                    self.lang_manager.get_text('error'),
                    self.lang_manager.get_text('unsupported_format')
                ))

            # Обновляем UI в основном потоке
            self.root.after(0, self.conversion_finished, success, output_path)

        except Exception as e:
            self.root.after(0, lambda: messagebox.showerror(
                self.lang_manager.get_text('error'), str(e)
            ))
            self.root.after(0, self.conversion_finished, False, None)

    def conversion_finished(self, success, output_path):
        self.progress_bar.set(1.0 if success else 0)
        self.convert_button.configure(state='normal')
        self.cancel_button.configure(state='disabled')

        if success:
            self.status_label.configure(text=self.lang_manager.get_text('success'))
            messagebox.showinfo(
                self.lang_manager.get_text('success'),
                f"{self.lang_manager.get_text('success')}\n{output_path}"
            )
        else:
            self.status_label.configure(text=self.lang_manager.get_text('error'))

        # Сбрасываем прогресс через 2 секунды
        self.root.after(2000, lambda: self.progress_bar.set(0))
        self.root.after(2000, lambda: self.status_label.configure(
            text=self.lang_manager.get_text('ready')
        ))

    def cancel_conversion(self):
        self.convert_button.configure(state='normal')
        self.cancel_button.configure(state='disabled')
        self.progress_bar.set(0)
        self.status_label.configure(text=self.lang_manager.get_text('ready'))

    def run(self):
        self.root.mainloop()


if __name__ == "__main__":
    app = QuickConverterApp()
    app.run()